#!/usr/bin/env python3
"""Build structured rules JSON for the SFL app.

Flag: parsed from the official DOCX (data/raw/flag.docx) — 1x1 tables are the top
categories, Heading 2 are the sections, and content tables (grade levels, field
dimensions, penalty tables) are captured as structured tables.
Tackle: parsed from data/raw/tackle.txt (unchanged Sophomore 6-man rules).

Output: rules.js (dual export for the browser + the serverless functions).
Both the lookup and the chat grounding derive from the same parsed content.
"""
import re, json, os
import docx
from docx.oxml.ns import qn

HERE = os.path.dirname(__file__)
RAW = os.path.join(HERE, "raw")
OUT = os.path.join(HERE, "..", "rules.js")
UPDATED = "As of August 26, 2026"

def norm(s): return re.sub(r"\s+", " ", (s or "").strip()).upper()
def slug(s): return re.sub(r"[^a-z0-9]+", "-", (s or "").lower()).strip("-")
def collapse(t): return re.sub(r"\n{3,}", "\n\n", (t or "")).strip()

TOP_CATEGORIES = {
    "GENERAL PROVISIONS", "SPORTSMANSHIP AND CONDUCT", "GENERAL GAME PROVISIONS",
    "GAME PLAY", "OFFENSE", "DEFENSE", "OFFICIALS AND PENALTIES",
    "SUMMARY FRESHMAN DIVISION PROVISIONS",
}

# ---------------- FLAG (docx) ----------------
def cells_of(t):
    return [[c.text.strip() for c in row.cells] for row in t.rows]

def all_same(row):
    vals = [c for c in row if c != ""]
    return len(set(vals)) == 1 and len(vals) >= 2

def build_ptable(cells):
    """3-col penalty table -> {notes:[...], rows:[{foul,definition,enforcement}]}"""
    notes, rows = [], []
    for i, row in enumerate(cells):
        if i == 0:
            continue  # header row (FOULS | DEFINITION | ENFORCEMENT)
        if all_same(row):
            if row[0] and row[0] not in notes:
                notes.append(row[0])
            continue
        foul = row[0].strip()
        if not foul:
            continue
        rows.append({
            "foul": foul,
            "definition": (row[1] if len(row) > 1 else "").strip(),
            "enforcement": (row[2] if len(row) > 2 else "").strip(),
        })
    return {"notes": notes, "rows": rows}

def build_grid(cells):
    """Generic reference table -> {headers, rows}."""
    return {"headers": cells[0], "rows": cells[1:]}

def parse_flag():
    d = docx.Document(os.path.join(RAW, "flag.docx"))
    pmap = {p._p: p for p in d.paragraphs}
    tmap = {t._tbl: t for t in d.tables}
    groups, cur_group, cur_sec, buf = [], None, None, []
    fresh_prov = []

    def flush():
        nonlocal buf, cur_sec
        if cur_sec is not None:
            cur_sec["text"] = collapse("\n".join(buf))
        buf = []

    for child in d.element.body.iterchildren():
        if child.tag == qn("w:p"):
            p = pmap.get(child)
            if p is None:
                continue
            txt = p.text.strip()
            if not txt:
                continue
            style = p.style.name if p.style else ""
            if style == "Heading 2":
                flush()
                cur_sec = {"title": txt, "id": slug(txt), "text": ""}
                (cur_group or {}).get("sections", []).append(cur_sec)
                continue
            buf.append(("• " + txt) if "List" in style else txt)
        elif child.tag == qn("w:tbl"):
            t = tmap.get(child)
            if t is None:
                continue
            cells = cells_of(t)
            # 1x1 = category divider (or TABLE OF CONTENTS)
            if len(cells) == 1 and len(cells[0]) == 1:
                label = norm(cells[0][0])
                if label == "TABLE OF CONTENTS":
                    continue
                flush()
                cur_group = {"title": label, "id": slug(label), "sections": []}
                groups.append(cur_group)
                cur_sec = None
                continue
            # content table
            title = norm(cur_sec["title"]) if cur_sec else ""
            if title in ("OFFENSIVE PENALTIES", "DEFENSIVE PENALTIES"):
                flush()
                cur_sec["ptable"] = build_ptable(cells)
            elif cur_group and cur_group["title"] == "SUMMARY FRESHMAN DIVISION PROVISIONS":
                for row in cells:
                    if len(row) >= 2 and row[0].strip():
                        fresh_prov.append({"label": row[0].strip(), "text": row[1].strip()})
            elif cur_sec is not None:
                flush()
                cur_sec.setdefault("tables", []).append(build_grid(cells))
    flush()
    # the freshman-summary category carries no sub-sections
    groups = [g for g in groups if g["sections"]]
    return groups, fresh_prov

def flag_fulltext(groups):
    out = []
    for g in groups:
        out.append(g["title"])
        for s in g["sections"]:
            out.append(s["title"])
            if s.get("text"):
                out.append(s["text"])
            for tbl in s.get("tables", []):
                out.append(" | ".join(tbl["headers"]))
                for r in tbl["rows"]:
                    out.append(" | ".join(r))
            pt = s.get("ptable")
            if pt:
                for n in pt["notes"]:
                    out.append(n)
                for r in pt["rows"]:
                    out.append(f"{r['foul']}: {r['definition']} — {r['enforcement']}")
    return collapse("\n\n".join(out))

# ---------------- TACKLE (txt) ----------------
TACKLE_HEADERS = [
    "GAME PLAY:", "The Clock will be stopped for the following reasons:",
    "SPECIAL TEAMS:", "POINT SYSTEM:", "SUBSTITUTION / PLAY RULE:",
    "REQUIRED and PROHIBITED EQUIPMENT:", "SPORTSMANSHIP: ZERO-TOLERANCE POLICY",
    "Safety/Weight Considerations:", "Field Dimensions:",
]
TACKLE_TITLES = {
    "GAME PLAY:": "Game Play",
    "The Clock will be stopped for the following reasons:": "Clock Stoppages",
    "SPECIAL TEAMS:": "Special Teams", "POINT SYSTEM:": "Point System",
    "SUBSTITUTION / PLAY RULE:": "Substitution / Play Rule",
    "REQUIRED and PROHIBITED EQUIPMENT:": "Required & Prohibited Equipment",
    "SPORTSMANSHIP: ZERO-TOLERANCE POLICY": "Sportsmanship: Zero-Tolerance Policy",
    "Safety/Weight Considerations:": "Safety / Weight Considerations",
    "Field Dimensions:": "Field Dimensions",
}

def parse_tackle():
    text = open(os.path.join(RAW, "tackle.txt")).read()
    hdr = {norm(h): h for h in TACKLE_HEADERS}
    sections, cur, buf = [], None, []
    def flush():
        nonlocal buf, cur
        if cur is not None:
            cur["text"] = collapse("\n".join(buf))
        buf = []
    for ln in text.split("\n"):
        ln = ln.replace("\f", "")
        s = ln.strip()
        if not s:
            buf.append("")
            continue
        if s == "2025 Season" or re.fullmatch(r"\d{1,3}", s):
            continue
        if norm(s) in hdr:
            flush()
            cur = {"title": TACKLE_TITLES[hdr[norm(s)]], "id": slug(TACKLE_TITLES[hdr[norm(s)]]), "text": ""}
            sections.append(cur)
            continue
        if cur is not None:
            buf.append(ln)
    flush()
    for sec in sections:
        if sec["id"] == "field-dimensions" and not sec["text"]:
            sec["text"] = ("Field dimensions for the Sophomore 6-man tackle field are set by the "
                           "league at the field each game day. Beyond the unique rules in this document, "
                           "game play is governed by TAPPS 6 Man and NCAA football rules.")
    groups = [{"title": "SOPHOMORE 6-MAN TACKLE RULES", "id": "tackle", "sections": sections}]
    full = []
    for sec in sections:
        full.append(sec["title"])
        if sec["text"]:
            full.append(sec["text"])
    return groups, collapse("\n\n".join(full))

def main():
    flag_groups, fresh_prov = parse_flag()
    flag_full = flag_fulltext(flag_groups)
    tackle_groups, tackle_full = parse_tackle()

    # Freshman override retained from an earlier league instruction (Jeff): no 4th down.
    # (Not in the official summary table — flagged separately for confirmation.)
    fresh_prov = fresh_prov + [{
        "label": "No 4th Down",
        "text": ("There is no 4th down in the Freshman Division. Because the offense does not get a "
                 "fourth down, a loss-of-down penalty (such as a false start) on third down ends the "
                 "offense's possession — effectively a turnover on downs."),
    }]
    prov_text = "FRESHMAN DIVISION PROVISIONS (adjustments to the flag rules)\n\n" + \
        "\n".join(f"- {p['label']}: {p['text']}" for p in fresh_prov)

    data = {
        "updated": UPDATED,
        "divisions": {
            "flag-older": {
                "id": "flag-older", "format": "Flag",
                "label": "Flag Football — Sophomore / Junior",
                "short": "Flag · Older divisions",
                "intro": "Standard SFL Flag Football rules for Sophomore (3rd–4th) and Junior (5th–6th) divisions.",
                "groups": flag_groups, "fullText": flag_full,
            },
            "flag-freshman": {
                "id": "flag-freshman", "format": "Flag",
                "label": "Flag Football — Freshman (1st–2nd grade)",
                "short": "Flag · Freshman",
                "intro": "SFL Flag Football rules with the Freshman Division adjustments applied. The provisions below override the standard rules where they differ.",
                "provisions": fresh_prov,
                "groups": flag_groups, "fullText": prov_text + "\n\n" + flag_full,
            },
            "tackle-sophomore": {
                "id": "tackle-sophomore", "format": "Tackle",
                "label": "Sophomore 6-Man Tackle",
                "short": "Tackle · Sophomore",
                "intro": "SFL Sophomore 6-Man Tackle Football rules (2025). Beyond the unique rules here, play is governed by TAPPS 6 Man and NCAA rules; conduct follows the SFL Flag Official Rulebook.",
                "groups": tackle_groups, "fullText": tackle_full,
            },
        },
        "order": ["flag-freshman", "flag-older", "tackle-sophomore"],
    }

    js = ("/* AUTO-GENERATED from the official SFL rulebooks by data/build_rules.py.\n"
          "   Do not edit by hand — re-run the builder when a rulebook changes. */\n"
          "const SFL_RULES = " + json.dumps(data, ensure_ascii=False, indent=2) + ";\n"
          "if (typeof module !== 'undefined' && module.exports) module.exports = SFL_RULES;\n"
          "if (typeof window !== 'undefined') window.SFL_RULES = SFL_RULES;\n")
    with open(OUT, "w") as f:
        f.write(js)

    print("FLAG groups:", len(flag_groups))
    for g in flag_groups:
        print("  -", g["title"], "->", [s["id"] for s in g["sections"]])
    pen = [s for g in flag_groups for s in g["sections"] if s.get("ptable")]
    print("penalty tables:", [(s["id"], len(s["ptable"]["rows"])) for s in pen])
    print("freshman provisions:", [p["label"] for p in fresh_prov])
    print("wrote", os.path.abspath(OUT))

if __name__ == "__main__":
    main()
